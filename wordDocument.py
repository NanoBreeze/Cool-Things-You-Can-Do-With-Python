'''
Copyright 2017 Linyi Cheng

Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated documentation files (the "Software"), to deal in the Software without restriction, including without limitation the rights to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies of the Software, and to permit persons to whom the Software is furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.
'''


from docx import Document
from docx.shared import RGBColor #setting text color
from docx.enum.text import WD_COLOR_INDEX #setting highlighter color
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #text alignment

'''
[event, eg birthday party] Invitation

Hey [name]!

Thanks for being a [positive adjective] friend! I would like to invite you to my
[event, eg birthday party] on [Day of week] [Month] [Day]. We will have a lot of fun together.

See you then!

Cheers,
[Name]
'''

if __name__ == '__main__':
    document = Document()

    event = input("What event are you hosting (eg, birthday party)? ")
    friend_name = input("What's your friend's name? ")
    adjective = input("Enter a nice adjective (eg, wonderful): ")
    day_of_week = input("What day of the week is your event on (eg, Monday)? ")
    month = input("What month is your event in? ")
    day = input("Which day is your event on? ")

    my_name = input("And most importantly, what's your name? ")


    heading = document.add_heading('{0} Invitation'.format(event.title()), level=0) #.title() capitalizes the first letter of each new word
    heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph('Hello {0}!'.format(friend_name))

    p = document.add_paragraph('Thanks for being a ')
    run = p.add_run(adjective + ' ')
    run.font.color.rgb = RGBColor(100, 100, 255)

    p.add_run('friend! I would like to invite you to my ')
    run = p.add_run(event + ' ')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW #cannot use RGBColor(...). Must use enumeration instead

    p.add_run('on ')

    run = p.add_run('{0} {1} {2}.'.format(day_of_week, day, month))
    run.bold = True
    p.add_run(' We will have a lot of fun together!')

    document.add_paragraph('See you then!')
    document.add_paragraph('Cheers,')
    document.add_paragraph(my_name)

    document.save('invitation.docx')

    print("Invitation generated!")